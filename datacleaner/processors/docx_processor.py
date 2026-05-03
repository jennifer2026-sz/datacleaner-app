"""Word document processor."""

from pathlib import Path


def process_docx(filepath: str | Path) -> str:
    """Extract text from a .docx file."""
    from docx import Document

    doc = Document(str(filepath))
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)
